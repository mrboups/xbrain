"""Pure document body extraction + chunking (Phase 24, DOCBODY-01).

`media.py` today embeds only `caption or filename`; the uploaded bytes go to MinIO and
the body is never extracted. This module is the reusable, SIDE-EFFECT-FREE core the upload
wiring (Plan 24-02) and the real-infra gate (Plan 24-03) both build against:

  - `extract_document(...)` dispatches by mime (PDF via pypdf, DOCX via python-docx,
    text/* + markdown decoded) and returns an `ExtractResult`. Unknown/binary mime is
    skipped cleanly (never a crash). A malformed/crafted document is caught HERE and
    turned into `skipped=True(reason=parse_error:...)` — it never propagates.
  - `chunk_text(...)` slices an extracted body into a bounded number of overlapping
    chunks, dropping empty/whitespace-only windows so no empty vector is ever embedded.

Purity is the contract (threat T-24-05): this module imports NO `app.config`, reads NO
team_scope, and makes NO network / DB / MinIO call. Every cap arrives as an EXPLICIT
keyword arg; the caller (doc_body_ingest.py) passes `settings.DOCBODY_*`. Heavy readers
(pypdf / python-docx) are imported LAZILY inside the per-surface helpers, matching the
embedders.py / neo4j / boto3 lazy-import convention already used across this app.
"""
from __future__ import annotations

import io
from dataclasses import dataclass

import structlog

log = structlog.get_logger(__name__)

# Exact mime strings Plan 02/03 dispatch on. DOCX is the OpenXML wordprocessing type.
PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@dataclass
class ExtractResult:
    """Outcome of a single extraction attempt.

    text          extracted (and possibly truncated) body; "" when skipped or no_text_layer
    no_text_layer a SUPPORTED surface (pdf/docx) yielded empty/near-empty text (scanned doc):
                  the parent item should be flagged, NOT given an empty-body chunk (D-24-03)
    skipped       nothing to embed — unknown mime, over-size, or a caught parse error
    surface       "pdf" | "docx" | "markdown" | "text" | None (None when skipped on mime/size)
    truncated     body exceeded max_total_chars and was cut to the cap before returning
    reason        machine-readable why for skipped/flagged outcomes (logs + parent metadata)
    """

    text: str
    no_text_layer: bool
    skipped: bool
    surface: str | None
    truncated: bool
    reason: str = ""


def _extract_pdf(data: bytes) -> str:
    """PDF text layer via pypdf (pure-Python, py3-none-any — no shell-out)."""
    from pypdf import PdfReader  # lazy — matches embedders/neo4j/boto3 convention

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _extract_docx(data: bytes) -> str:
    """DOCX paragraph + table-cell text via python-docx (lazy import)."""
    import docx  # lazy — python-docx pulls lxml (compiled); import only when needed

    document = docx.Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in document.paragraphs]
    # Table cells carry body text too (e.g. a spec laid out in a table). Include them so a
    # table-only document is not falsely flagged no_text_layer.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_text(data: bytes) -> str:
    """Decode text/* + markdown. errors='replace' — hostile bytes can never execute
    or raise here (threat T-24-04); the content is inert text, never rendered."""
    return data.decode("utf-8", errors="replace")


def extract_document(
    data: bytes,
    mime: str,
    filename: str | None,
    *,
    max_file_bytes: int,
    max_total_chars: int,
    min_chars: int,
) -> ExtractResult:
    """Extract the text body of an uploaded document, dispatching by mime.

    Order is deliberate:
      1. SIZE GUARD FIRST — an over-cap file is skipped before any parser touches it
         (never parse attacker bytes past the cap; DoS defence T-24-02).
      2. dispatch by mime; unknown/binary -> skipped, surface=None (no crash).
      3. per-surface extraction wrapped in try/except -> a malformed doc becomes
         skipped(reason=parse_error:<type>), never an exception (T-24-01).
      4. TRUNCATE to max_total_chars (token-budget guard) before returning.
      5. NO-TEXT-LAYER — for pdf/docx only, a body under min_chars means "scanned /
         image-only": flag it and emit NO text (so no empty chunk downstream, D-24-03).
         For text/markdown, short is just short.
    """
    # 1. size guard — before any parse
    if len(data) > max_file_bytes:
        log.info("doc_extract.skip_too_large", filename=filename, size=len(data), cap=max_file_bytes)
        return ExtractResult(
            text="", no_text_layer=False, skipped=True, surface=None, truncated=False,
            reason="file_too_large",
        )

    # 2. dispatch
    if mime == PDF_MIME:
        surface = "pdf"
        extractor = _extract_pdf
    elif mime == DOCX_MIME:
        surface = "docx"
        extractor = _extract_docx
    elif mime.startswith("text/"):
        # "text/markdown" / "text/x-markdown" keep their markdown surface label; every
        # other text/* is plain text. Both go through the same utf-8 decode.
        surface = "markdown" if "markdown" in mime else "text"
        extractor = _extract_text
    else:
        # unknown / binary mime — nothing to extract, and definitely no crash
        return ExtractResult(
            text="", no_text_layer=False, skipped=True, surface=None, truncated=False,
            reason="unsupported_mime",
        )

    # 3. per-surface extraction, fail-soft
    try:
        text = extractor(data)
    except Exception as exc:  # broad on purpose — a crafted doc must NEVER propagate
        log.warning(
            "doc_extract.parse_error", filename=filename, surface=surface, error=str(exc)
        )
        return ExtractResult(
            text="", no_text_layer=False, skipped=True, surface=surface, truncated=False,
            reason=f"parse_error:{type(exc).__name__}",
        )

    # 4. truncate (token-budget guard)
    truncated = False
    if len(text) > max_total_chars:
        text = text[:max_total_chars]
        truncated = True

    # 5. no_text_layer — supported binary surfaces only
    if surface in ("pdf", "docx") and len(text.strip()) < min_chars:
        return ExtractResult(
            text="", no_text_layer=True, skipped=False, surface=surface, truncated=truncated,
            reason="no_text_layer",
        )

    return ExtractResult(
        text=text, no_text_layer=False, skipped=False, surface=surface, truncated=truncated,
    )


def chunk_text(
    text: str,
    *,
    chunk_size: int,
    overlap: int,
    max_chunks: int,
) -> list[str]:
    """Slice `text` into overlapping character windows.

    - sliding window of `chunk_size` chars, advancing by `chunk_size - overlap`
    - stop at `max_chunks` (chunk-count cap, T-24-02) even if text remains
    - strip each window and DROP empty/whitespace-only ones (never embed an empty vector)
    - returns [] for empty/whitespace-only input
    """
    if not text or not text.strip():
        return []

    # Guard against a pathological overlap >= chunk_size (would never advance / infinite loop).
    step = chunk_size - overlap
    if step <= 0:
        step = chunk_size

    chunks: list[str] = []
    start = 0
    length = len(text)
    while start < length and len(chunks) < max_chunks:
        window = text[start : start + chunk_size].strip()
        if window:
            chunks.append(window)
        start += step
    return chunks
